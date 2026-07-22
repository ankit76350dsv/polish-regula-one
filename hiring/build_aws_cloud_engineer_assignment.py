from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_DOCX = OUTPUT_DIR / "AWS_Cloud_Engineer_Practical_Assignment_CloudVault.docx"
DIAGRAM_PATH = OUTPUT_DIR / "cloudvault_architecture.png"

# Resolved design system: decision_memo alias of standard_business_brief.
# Page: Letter, 1-inch margins, 0.492-inch header/footer, 9360 DXA content width.
# Typography: Arial 11 pt, 6 pt after, 1.10 line spacing.
# H1: 16 pt / #2E74B5 / 12 pt before / 6 pt after.
# H2: 13 pt / #2E74B5 / 10 pt before / 5 pt after.
# H3: 12 pt / #1F4D78 / 8 pt before / 4 pt after.
# Lists: marker 0.25 in, text 0.5 in, hanging 0.25 in, 8 pt after, 1.167 spacing.
# Tables: 9360 DXA, indent 120 DXA, margins 80/80/120/120, #F2F4F7 header.

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "172B4D"
INK = "1D2733"
MUTED = "5B6775"
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
ORANGE = "FF9900"
GREEN = "247A4A"
RED = "A33A3A"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=11, bold=False, italic=False, color=INK, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D5DAE0", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, received {sum(widths_dxa)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_borders(table)


def set_paragraph_border_bottom(paragraph, color=ORANGE, size=18, space=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 12, 6),
        "Heading 2": (13, BLUE, 10, 5),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("AWS CLOUD ENGINEER PRACTICAL ASSIGNMENT  |  CANDIDATE BRIEF")
    set_run_font(r, size=8.5, bold=True, color=MUTED)
    set_paragraph_border_bottom(p, color="D8DDE3", size=6, space=4)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("CloudVault  •  Page ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(p)


def create_numbering(doc, fmt, text, left=720, hanging=360):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list(doc, items, numbered=False):
    num_id = create_numbering(doc, "decimal" if numbered else "bullet", "%1." if numbered else "•")
    for item in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_el)
        p_pr.insert(0, num_pr)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        set_run_font(r)
    return p


def add_paragraph(doc, text="", bold_lead=None, color=INK, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_callout(doc, label, text, fill=CALLOUT, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.line_spacing = 0.5
    return table


def add_table(doc, headers, rows, widths_dxa, alignments=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        if alignments:
            p.alignment = alignments[idx]
        r = p.add_run(value)
        set_run_font(r, size=9.5, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if alignments:
                p.alignment = alignments[idx]
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    after.paragraph_format.line_spacing = 0.5
    return table


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_shading(table.cell(0, 0), "F7F8FA")
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(text.strip("\n").splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, size=9, color="263238", name="Courier New")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.line_spacing = 0.5
    return table


def load_font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def draw_centered_multiline(draw, box, text, font, fill, spacing=4):
    x1, y1, x2, y2 = box
    max_chars = max(10, int((x2 - x1) / (font.size * 0.55))) if hasattr(font, "size") else 20
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(wrap(paragraph, max_chars) or [""])
    bbox = draw.multiline_textbbox((0, 0), "\n".join(lines), font=font, spacing=spacing, align="center")
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - tw) / 2, (y1 + y2 - th) / 2), "\n".join(lines), font=font, fill=fill, spacing=spacing, align="center")


def draw_box(draw, box, title, subtitle="", fill="#FFFFFF", outline="#7D8A99", title_color="#172B4D"):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    title_font = load_font(25, bold=True)
    sub_font = load_font(18, bold=False)
    if subtitle:
        draw_centered_multiline(draw, (x1 + 12, y1 + 10, x2 - 12, y1 + 58), title, title_font, title_color)
        draw_centered_multiline(draw, (x1 + 12, y1 + 54, x2 - 12, y2 - 8), subtitle, sub_font, "#4C5B6B")
    else:
        draw_centered_multiline(draw, box, title, title_font, title_color)


def arrow(draw, start, end, color="#526579", width=4):
    draw.line([start, end], fill=color, width=width)
    import math
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 15
    spread = 0.55
    p1 = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
    p2 = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
    draw.polygon([end, p1, p2], fill=color)


def build_architecture_diagram():
    image = Image.new("RGB", (1600, 1000), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    draw.text((70, 34), "CloudVault reference flow", font=title_font, fill="#172B4D")
    draw.line((70, 86, 1530, 86), fill="#FF9900", width=6)

    draw_box(draw, (70, 145, 300, 255), "Client / Operator", "Upload and status requests", fill="#FFF7E8", outline="#E09B2D")
    draw_box(draw, (400, 135, 690, 265), "WAF + ALB", "Rate control and application routing", fill="#EEF5FB", outline="#2E74B5")
    draw_box(draw, (800, 125, 1160, 275), "EC2 Auto Scaling", "Management API in private subnets\nIAM role • IMDSv2 • SSM", fill="#EEF5FB", outline="#2E74B5")
    draw_box(draw, (1270, 145, 1530, 255), "DynamoDB", "Job status and idempotency", fill="#F3EEFB", outline="#7057A8")

    draw_box(draw, (120, 390, 400, 520), "S3 Intake", "Private • versioned • SSE-KMS", fill="#F0F8EF", outline="#4C8C43")
    draw_box(draw, (500, 390, 790, 520), "SQS + DLQ", "Buffering • retries • redrive", fill="#F9F0F6", outline="#A75486")
    draw_box(draw, (890, 380, 1190, 530), "Lambda Processor", "Validate • checksum • isolate tenant\nstructured logs • partial failures", fill="#FFF3E8", outline="#D47B24")
    draw_box(draw, (1290, 375, 1530, 535), "S3 Results", "Processed and quarantine\nSSE-KMS", fill="#F0F8EF", outline="#4C8C43")

    draw_box(draw, (160, 665, 430, 785), "EventBridge", "Daily operational schedule", fill="#F4F1FB", outline="#7057A8")
    draw_box(draw, (530, 655, 830, 795), "Reporting Lambda", "Metrics summary and report", fill="#FFF3E8", outline="#D47B24")
    draw_box(draw, (930, 665, 1190, 785), "SNS", "Completion and failure alerts", fill="#F9F0F6", outline="#A75486")
    draw_box(draw, (1290, 655, 1530, 795), "Reports S3", "Encrypted daily reports", fill="#F0F8EF", outline="#4C8C43")

    arrow(draw, (300, 200), (400, 200))
    arrow(draw, (690, 200), (800, 200))
    arrow(draw, (1160, 200), (1270, 200))
    arrow(draw, (900, 275), (385, 390))
    arrow(draw, (400, 455), (500, 455))
    arrow(draw, (790, 455), (890, 455))
    arrow(draw, (1190, 455), (1290, 455))
    arrow(draw, (1090, 380), (1330, 255))
    arrow(draw, (430, 725), (530, 725))
    arrow(draw, (830, 725), (930, 725))
    arrow(draw, (830, 745), (1290, 745))

    draw.rounded_rectangle((70, 875, 1530, 955), radius=14, fill="#F4F6F9", outline="#AEB8C2", width=2)
    footer_font = load_font(23, bold=True)
    draw_centered_multiline(
        draw,
        (85, 885, 1515, 945),
        "Cross-cutting controls: IAM least privilege  •  KMS  •  Secrets Manager  •  CloudWatch  •  CloudTrail  •  Infrastructure as Code",
        footer_font,
        "#33475B",
    )
    image.save(DIAGRAM_PATH, dpi=(180, 180))


def add_picture_with_alt(doc, path, alt_text, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(6.45))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", "CloudVault architecture diagram")
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(58)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("CANDIDATE PRACTICAL ASSIGNMENT")
    set_run_font(r, size=11, bold=True, color=ORANGE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("CloudVault")
    set_run_font(r, size=32, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("Secure Multi-Tenant Document Processing Platform")
    set_run_font(r, size=15, color=DARK_BLUE)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(22)
    set_paragraph_border_bottom(rule, color=ORANGE, size=20, space=4)

    facts = [
        ("Target level", "AWS Cloud Engineer • approximately 2 years of experience"),
        ("Expected effort", "12–16 hours"),
        ("Submission window", "4–5 calendar days"),
        ("Primary focus", "Architecture, security, automation, reliability and operations"),
        ("Version", "1.0 • Candidate brief"),
    ]
    for label, value in facts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.5, bold=True, color=NAVY)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=MUTED)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(24)
    add_callout(
        doc,
        "Purpose",
        "Design and deploy a production-minded AWS workload. Application code may be intentionally small; the assessment emphasizes sound cloud engineering decisions and evidence that the system works.",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Independent hiring exercise • Not connected to any existing company project")
    set_run_font(r, size=9, italic=True, color=MUTED)
    doc.add_page_break()


def build_document():
    build_architecture_diagram()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    props = doc.core_properties
    props.title = "AWS Cloud Engineer Practical Assignment — CloudVault"
    props.subject = "Candidate take-home assessment for an AWS Cloud Engineer"
    props.author = "Hiring Team"
    props.keywords = "AWS, EC2, Lambda, S3, KMS, IAM, SQS, DynamoDB, CloudWatch, Terraform"
    props.comments = "Candidate-facing assessment brief"

    add_cover(doc)

    add_heading(doc, "1. Assignment overview", 1)
    add_paragraph(
        doc,
        "Build CloudVault, a secure AWS platform that accepts JSON or PDF documents, processes them asynchronously, tracks every job, isolates tenant data, and exposes enough operational evidence to investigate failures. The system must be deployed from code and must not depend on manually created production resources.",
    )
    add_callout(
        doc,
        "Scope boundary",
        "A polished frontend is not required. A minimal API and simple processor are sufficient. Spend your time on AWS design, security, automation, monitoring, recoverability and clear engineering reasoning.",
    )

    add_heading(doc, "1.1 Engagement parameters", 2)
    add_table(
        doc,
        ["Item", "Requirement"],
        [
            ("Expected effort", "12–16 hours; document incomplete items instead of exceeding the time box."),
            ("Submission window", "4–5 calendar days from receipt."),
            ("AWS account", "Use a personal sandbox or an account supplied by the hiring team."),
            ("Region", "Deploy to one AWS Region. Multi-region implementation is not required."),
            ("Infrastructure", "Terraform, CloudFormation, AWS CDK or SAM. Console-only builds are not accepted."),
            ("Application language", "Python, Node.js, Java, Go or another reasonable language."),
            ("Evidence", "Provide deployment output, test evidence and screenshots or exported metrics as appropriate."),
        ],
        [2000, 7360],
    )

    add_heading(doc, "1.2 Success outcome", 2)
    add_list(
        doc,
        [
            "A valid tenant document can be uploaded without making an S3 bucket public.",
            "The document is processed asynchronously and reaches a terminal status.",
            "Invalid documents are quarantined without being confused with technical failures.",
            "Temporary failures are retried and unrecoverable messages reach a dead-letter queue.",
            "Duplicate events do not create duplicate business outcomes.",
            "An EC2 failure is detected and the service recovers through Auto Scaling.",
            "Logs, metrics and audit events allow an operator to explain what happened to a job.",
            "All resources can be recreated and removed using documented commands.",
        ],
    )

    add_heading(doc, "2. Business scenario", 1)
    add_paragraph(
        doc,
        "A software company receives business documents from multiple customers. Each customer is represented as a tenant. Uploaded content may contain sensitive information, so storage must be private and encrypted. Processing demand is uneven: the platform may receive only a few files during normal operation and a sudden burst during a customer import.",
    )
    add_paragraph(
        doc,
        "The operations team needs an API for upload initiation and status lookup, automatic processing, alerts for failure, a daily summary, and evidence suitable for troubleshooting or a security review. Tenant A must never be able to read Tenant B’s job or document.",
    )

    add_heading(doc, "2.1 Mandatory AWS services", 2)
    add_table(
        doc,
        ["Capability", "Required AWS services"],
        [
            ("Network and compute", "Amazon VPC, EC2, EC2 Auto Scaling, Application Load Balancer, Systems Manager"),
            ("Storage and processing", "Amazon S3, AWS Lambda, Amazon SQS and an SQS dead-letter queue"),
            ("State and scheduling", "Amazon DynamoDB, Amazon EventBridge"),
            ("Security", "AWS IAM, AWS KMS, AWS Secrets Manager, AWS WAF"),
            ("Operations and audit", "Amazon CloudWatch, AWS CloudTrail, Amazon SNS"),
            ("Delivery", "Infrastructure as Code and a CI/CD pipeline"),
        ],
        [2600, 6760],
    )

    add_heading(doc, "2.2 Assessment emphasis", 2)
    add_list(
        doc,
        [
            "Translate a business requirement into a coherent AWS design and explain meaningful trade-offs.",
            "Apply IAM, KMS, network and data-access controls without relying on broad administrative permissions.",
            "Build event-driven processing that remains correct during retries, duplicates and partial failures.",
            "Create repeatable delivery automation and enough telemetry to diagnose an operational problem.",
            "Balance availability and security with the cost constraints of a short-lived assessment account.",
        ],
    )
    add_callout(
        doc,
        "Not required",
        "A polished user interface, complex document parsing, a multi-region deployment, enterprise-scale test data or perfect completion of every optional improvement. Clearly document any deliberate simplification.",
    )

    add_heading(doc, "3. Reference architecture and workflow", 1)
    add_picture_with_alt(
        doc,
        DIAGRAM_PATH,
        "Architecture flow from client through WAF and an Application Load Balancer to an EC2 Auto Scaling group. Documents are uploaded to S3, queued in SQS, processed by Lambda, recorded in DynamoDB, written to processed or quarantine S3 storage, and reported through EventBridge, Lambda and SNS. IAM, KMS, Secrets Manager, CloudWatch, CloudTrail and Infrastructure as Code are cross-cutting controls.",
        "Figure 1. Reference flow. Equivalent architectures are acceptable when trade-offs are explained.",
    )
    add_paragraph(
        doc,
        "You may refine this architecture. Any change should preserve the required services and business outcomes and should be justified in the README.",
        italic=True,
        color=MUTED,
    )

    add_heading(doc, "3.1 End-to-end processing sequence", 2)
    add_list(
        doc,
        [
            "A client calls the EC2-hosted management API to create a job and request a pre-signed S3 upload URL.",
            "The API creates a PENDING job in DynamoDB and returns a time-limited URL tied to the expected tenant and object key.",
            "The client uploads a JSON or PDF document directly to the private S3 intake location.",
            "S3 delivers an object-created event to the processing SQS queue.",
            "Lambda consumes the message, validates the object and metadata, computes a SHA-256 checksum, and performs an idempotency check.",
            "Valid documents are written to processed storage; invalid business documents are written to quarantine storage with a safe rejection reason.",
            "DynamoDB reaches COMPLETED, REJECTED or FAILED, and SNS publishes the appropriate operational notification.",
            "EventBridge invokes a reporting Lambda on a schedule and stores an encrypted operational report in S3.",
        ],
        numbered=True,
    )

    add_heading(doc, "4. Functional requirements", 1)
    add_heading(doc, "4.1 EC2 management API", 2)
    add_table(
        doc,
        ["Method and path", "Required behavior"],
        [
            ("GET /health", "Return process health and dependency-readiness information suitable for ALB health checks."),
            ("POST /api/v1/documents/upload-url", "Validate request metadata, create a PENDING job and return a short-lived pre-signed upload URL."),
            ("GET /api/v1/documents/{jobId}", "Return the job only when it belongs to the authenticated or asserted tenant."),
            ("GET /api/v1/documents", "Return a tenant-scoped, paginated list without using a DynamoDB table scan."),
            ("GET /api/v1/operations/summary", "Return basic counts for completed, rejected and failed jobs."),
        ],
        [3000, 6360],
    )
    add_list(
        doc,
        [
            "Use versioned routes under /api/v1.",
            "Return consistent JSON error responses and appropriate HTTP status codes.",
            "Do not return internal stack traces, AWS credentials, raw secrets or another tenant’s identifiers.",
            "A production authentication system is not required. Use a signed JWT, an API-key-to-tenant mapping, or a clearly documented test identity mechanism and explain the production replacement.",
        ],
    )

    add_heading(doc, "4.2 Accepted document and metadata rules", 2)
    add_list(
        doc,
        [
            "Accept JSON and PDF only; verify extension and content type and document any stronger content validation used.",
            "Default maximum object size: 10 MB. Make the limit configurable.",
            "Require jobId, tenantId, customerId, documentType, fileName and createdAt metadata.",
            "Reject malformed JSON metadata and invalid identifiers without logging the document body.",
            "Calculate and persist a SHA-256 checksum for accepted content.",
            "Use deterministic output keys so replay does not create duplicate processed objects.",
        ],
    )

    add_heading(doc, "4.3 Job state model", 2)
    add_code_block(
        doc,
        """PENDING -> PROCESSING -> COMPLETED
                     |-> REJECTED   (business validation failure)
                     |-> FAILED     (terminal technical failure)

Allowed terminal states: COMPLETED, REJECTED, FAILED""",
    )
    add_paragraph(
        doc,
        "Use conditional updates or an equivalent concurrency control so that late, duplicate or replayed events cannot incorrectly move a terminal job backward.",
    )

    add_heading(doc, "5. Platform requirements", 1)
    add_heading(doc, "5.1 VPC, load balancing and EC2", 2)
    add_list(
        doc,
        [
            "Create a custom VPC spanning at least two Availability Zones.",
            "Place the ALB in public subnets and EC2 application instances in private subnets.",
            "Use a launch template and Auto Scaling group. Production design must support at least two instances; a one-instance development setting is acceptable to control assessment cost.",
            "Configure an ALB health check and one target-tracking scaling policy.",
            "Do not assign public IP addresses to EC2 and do not expose SSH. Use Systems Manager Session Manager for administration.",
            "Require IMDSv2 and use an EC2 instance profile instead of access keys.",
            "Allow the application port only from the ALB security group.",
            "Run the API as a non-root service that starts on boot and restarts after a crash.",
            "Explain the choice between NAT Gateway, VPC endpoints or a combined design, including cost implications.",
        ],
    )
    add_callout(
        doc,
        "TLS during assessment",
        "If you do not own a domain, HTTP is acceptable for a short-lived sandbox demonstration. Your production design must show ACM-managed TLS termination, secure redirection and an appropriate DNS approach.",
        fill="FFF8E8",
        accent="8A5A00",
    )

    add_heading(doc, "5.2 S3 and document lifecycle", 2)
    add_list(
        doc,
        [
            "Provide logical storage for intake, processed, quarantine and audit/report data. Separate buckets or clearly isolated prefixes are acceptable when justified.",
            "Enable Block Public Access, versioning and customer-managed KMS encryption.",
            "Deny non-TLS requests through bucket policy and restrict workload roles to required prefixes and operations.",
            "Use tenant-scoped object keys and avoid predictable access to another tenant’s content.",
            "Configure lifecycle rules and explain retention, deletion and recovery behavior.",
            "Demonstrate that an unauthenticated or unauthorized public object request is denied.",
        ],
    )

    add_heading(doc, "5.3 AWS KMS", 2)
    add_list(
        doc,
        [
            "Create customer-managed symmetric keys for application documents and audit logs, with separate aliases.",
            "Enable automatic rotation and separate key administration from workload use.",
            "Grant decrypt or data-key permissions only to the roles and services that require them.",
            "Avoid broad principals and explain every wildcard retained in a key policy.",
            "Document the operational response to a disabled key and the safeguards around scheduled deletion.",
        ],
    )

    add_heading(doc, "5.4 SQS and Lambda processing", 2)
    add_list(
        doc,
        [
            "Use an encrypted main queue and encrypted dead-letter queue with a documented redrive policy.",
            "Set the visibility timeout in relation to the Lambda timeout and explain the chosen values.",
            "Consume messages in batches and implement partial batch failure reporting.",
            "Use environment variables for resource identifiers and a dedicated Lambda execution role.",
            "Configure deliberate memory, timeout, batching and concurrency values.",
            "Use structured logs containing jobId, tenantId, outcome and duration; never log the document body.",
            "Separate a business rejection from a retryable technical error.",
            "Describe and demonstrate a safe DLQ replay procedure.",
        ],
    )

    add_heading(doc, "5.5 DynamoDB", 2)
    add_list(
        doc,
        [
            "Model tenant-scoped status and list queries without relying on Scan for normal API operations.",
            "Store job state, source and result keys, checksum, timestamps, failure reason and attempt count.",
            "Use conditional writes or transactions for idempotency and safe state transitions.",
            "Enable point-in-time recovery and customer-managed KMS encryption.",
            "Use TTL only where the business retention decision supports deletion.",
            "Document partition keys, sort keys, secondary indexes and expected access patterns.",
        ],
    )

    add_heading(doc, "5.6 EventBridge, reporting and SNS", 2)
    add_list(
        doc,
        [
            "Use EventBridge to invoke a daily reporting Lambda. A five-minute schedule may be used during demonstration.",
            "Report received, completed, rejected and failed counts; DLQ depth; and average processing duration.",
            "Store the report in encrypted S3 with a predictable date-based key.",
            "Publish SNS notifications for completion, rejection and significant operational failures.",
            "An email subscription may remain unconfirmed if you prefer not to provide a personal address.",
        ],
    )

    add_heading(doc, "5.7 Secrets Manager and configuration", 2)
    add_list(
        doc,
        [
            "Store at least one realistic secret such as an application signing secret or external webhook token in Secrets Manager.",
            "Do not place the secret value in source code, EC2 user data, CI logs or committed variable files.",
            "Retrieve the secret through a workload IAM role and describe rotation and cache behavior.",
            "Keep non-secret configuration in environment variables, Parameter Store or another justified source.",
        ],
    )

    add_heading(doc, "5.8 WAF, IAM and tenant isolation", 2)
    add_list(
        doc,
        [
            "Attach a WAF web ACL to the ALB with an AWS-managed common rule group, a rate-based rule and a request-size restriction.",
            "Use separate IAM roles for EC2, document processing, reporting and delivery automation.",
            "Do not use workload IAM users, long-lived access keys or AdministratorAccess.",
            "Avoid Action: * and Resource: * where resource-level permissions are available; justify unavoidable exceptions.",
            "Enforce tenant scope in API logic, DynamoDB access patterns and S3 object keys.",
            "Explain how WAF false positives and IAM authorization failures would be investigated.",
        ],
    )

    add_heading(doc, "6. Observability, audit and recovery", 1)
    add_heading(doc, "6.1 CloudWatch requirements", 2)
    add_table(
        doc,
        ["Area", "Minimum telemetry"],
        [
            ("ALB", "Request count, target response time, 4xx/5xx and healthy target count"),
            ("EC2", "CPU, status checks and application logs"),
            ("Lambda", "Invocations, duration, errors and throttles"),
            ("SQS", "Visible messages, age of oldest message and DLQ depth"),
            ("Application", "At least one custom metric for successful or rejected processing"),
        ],
        [2200, 7160],
    )
    add_list(
        doc,
        [
            "Create a single CloudWatch dashboard covering the minimum telemetry.",
            "Configure alarms for no healthy targets, EC2 status failure, Lambda errors or throttles, queue backlog, DLQ messages and elevated ALB 5xx responses.",
            "Set explicit retention on every application-created log group.",
            "Make one job traceable across EC2, S3 event processing and Lambda by jobId.",
        ],
    )

    add_heading(doc, "6.2 CloudTrail and audit evidence", 2)
    add_list(
        doc,
        [
            "Create a trail for management events, write it to a dedicated encrypted S3 bucket and enable log-file validation.",
            "Restrict CloudTrail bucket and KMS access to intended principals and services.",
            "Enable relevant S3 data events or explain the cost decision and production recommendation.",
            "Document how you would investigate a deleted object, disabled KMS key, changed security-group rule and changed IAM policy.",
        ],
    )

    add_heading(doc, "6.3 Backup and recovery", 2)
    add_list(
        doc,
        [
            "Use S3 versioning and DynamoDB point-in-time recovery.",
            "Define an RPO and RTO appropriate for this sandbox workload and a more realistic production target.",
            "Explain how infrastructure, configuration and data would be restored after accidental deletion.",
            "Describe—not implement—a regional disaster-recovery approach and identify which data requires replication.",
        ],
    )

    add_heading(doc, "7. Infrastructure as Code and CI/CD", 1)
    add_heading(doc, "7.1 Infrastructure as Code", 2)
    add_list(
        doc,
        [
            "Create all important resources from code. Small verification actions in the console are acceptable; manual infrastructure dependencies are not.",
            "Use reusable modules, constructs or nested stacks for networking, compute, storage, IAM/KMS, processing, monitoring and audit concerns.",
            "Separate environment-specific inputs, validate important variables and use consistent resource tags.",
            "Return useful outputs such as the ALB endpoint, bucket names, queue URLs and dashboard name.",
            "Do not commit Terraform state, generated deployment credentials, secrets or .env files.",
            "If using Terraform, explain how remote state, locking, encryption and state access would be secured.",
        ],
    )

    add_heading(doc, "7.2 CI/CD", 2)
    add_list(
        doc,
        [
            "Provide a pipeline for source checkout, formatting, IaC validation, static security checks, application tests, packaging and deployment to a development environment.",
            "Prefer OIDC or another short-lived credential mechanism for CI access to AWS.",
            "Do not store static AWS access keys in repository secrets unless you explicitly document why and how they would be replaced.",
            "A production approval stage is optional; describe the controls you would add for production.",
        ],
    )

    add_heading(doc, "8. Required verification scenarios", 1)
    add_paragraph(
        doc,
        "Provide an automated test, command transcript, screenshot or concise observation for each scenario. Clearly distinguish tests you executed from production recommendations you only documented.",
    )
    verification_rows = [
        ("1", "Valid JSON upload", "Job reaches COMPLETED and encrypted result exists."),
        ("2", "Valid PDF upload", "Supported content completes without being logged."),
        ("3", "Unsupported extension", "Job reaches REJECTED and object is quarantined."),
        ("4", "Oversized object", "Configurable limit is enforced safely."),
        ("5", "Malformed metadata", "Request or processing is rejected with a safe reason."),
        ("6", "Missing tenant ID", "No unscoped job or object is created."),
        ("7", "Duplicate event", "Only one business result is produced."),
        ("8", "Temporary Lambda failure", "Message is retried without corrupting job state."),
        ("9", "Repeated technical failure", "Message reaches DLQ and alarm evidence is visible."),
        ("10", "Cross-tenant read", "API denies access and does not reveal object metadata."),
        ("11", "EC2 termination", "Auto Scaling launches a replacement and health recovers."),
        ("12", "Public S3 request", "Anonymous or unauthorized access is denied."),
        ("13", "API request burst", "WAF rate rule or documented test threshold responds."),
        ("14", "Cleanup", "Destroy process removes assessment resources as documented."),
    ]
    add_table(
        doc,
        ["#", "Scenario", "Expected evidence"],
        verification_rows,
        [600, 3000, 5760],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=9,
    )

    add_heading(doc, "9. Deliverables and repository structure", 1)
    add_code_block(
        doc,
        """cloudvault/
├── infrastructure/
├── ec2-application/
├── lambda/
│   ├── document-processor/
│   └── daily-reporter/
├── tests/
├── diagrams/
├── scripts/
├── .github/workflows/ or pipeline/
├── README.md
├── SECURITY.md
├── RUNBOOK.md
└── COST-ESTIMATE.md""",
    )
    add_heading(doc, "9.1 README expectations", 2)
    add_list(
        doc,
        [
            "Architecture diagram and end-to-end data flow.",
            "Prerequisites and exact deploy, verify and destroy commands.",
            "IAM, KMS, network and tenant-isolation decisions.",
            "DynamoDB access patterns and idempotency strategy.",
            "Failure handling, DLQ replay and operational runbook links.",
            "Estimated monthly production cost and actual assessment cost where available.",
            "Assumptions, known limitations, incomplete items and next production improvements.",
        ],
    )

    add_heading(doc, "9.2 Submission package", 2)
    add_list(
        doc,
        [
            "A private repository or archive containing source code and documentation.",
            "A short evidence folder or document showing successful and failure-path verification.",
            "A concise presentation or README section suitable for a 20-minute technical walkthrough.",
            "No credentials, secret values, private keys, state files or customer information.",
        ],
    )

    add_heading(doc, "10. Demonstration acceptance checklist", 1)
    add_list(
        doc,
        [
            "Deploy the infrastructure using the documented process.",
            "Show EC2 targets registered behind the ALB and explain security-group boundaries.",
            "Create an upload job and upload a valid document using the pre-signed URL.",
            "Show the job status, encrypted processed object and matching DynamoDB record.",
            "Show invalid-document quarantine and a safe rejection reason.",
            "Show duplicate-event behavior, Lambda retry behavior and a message reaching the DLQ.",
            "Locate one job in structured logs and show the relevant dashboard or alarm.",
            "Find an infrastructure change or object action in CloudTrail.",
            "Terminate an EC2 instance and show Auto Scaling replacement behavior.",
            "Demonstrate that public S3 access and cross-tenant status access are denied.",
            "Run or explain the complete infrastructure cleanup process.",
        ],
        numbered=True,
    )

    add_heading(doc, "11. Evaluation rubric", 1)
    rubric_rows = [
        ("Architecture and service selection", "10", "Coherent data flow, sound trade-offs and limited unnecessary complexity."),
        ("VPC, ALB, EC2 and Auto Scaling", "15", "Private compute, secure access, health checks, bootstrapping and recovery."),
        ("Lambda, SQS and event processing", "15", "Retries, DLQ, partial failure, idempotency and correct error classification."),
        ("IAM, KMS, S3 and tenant security", "20", "Least privilege, private encryption, safe key policy and tenant isolation."),
        ("Infrastructure as Code", "15", "Reproducible, modular, validated and cleanly destroyable infrastructure."),
        ("Monitoring, audit and recovery", "10", "Useful telemetry, alarms, CloudTrail evidence and recovery reasoning."),
        ("CI/CD and automation", "5", "Safe short-lived authentication and meaningful validation stages."),
        ("Testing and failure handling", "5", "Credible success, rejection, retry, replay and isolation tests."),
        ("Documentation and communication", "5", "Clear instructions, decisions, limitations and technical explanation."),
    ]
    add_table(
        doc,
        ["Evaluation area", "Points", "What strong evidence looks like"],
        rubric_rows,
        [3000, 900, 5460],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=9,
    )
    add_callout(
        doc,
        "Recommended pass threshold",
        "65/100 overall, including at least 12/20 in IAM, KMS, S3 and tenant security. A strong explanation of an intentional trade-off may receive credit even when the implementation is simplified for time or cost.",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )

    add_heading(doc, "11.1 Critical security failures", 2)
    add_paragraph(doc, "A submission may be rejected regardless of total score when it includes any of the following:")
    add_list(
        doc,
        [
            "Committed AWS credentials, private keys or real secret values.",
            "Public access to document buckets or deliberate exposure of tenant content.",
            "AdministratorAccess attached to an application workload.",
            "SSH exposed to 0.0.0.0/0 or public EC2 access without a defensible constraint.",
            "Unencrypted document storage or hardcoded credentials on EC2 or Lambda.",
            "No meaningful tenant isolation, retry handling or Infrastructure as Code.",
            "A deployment that cannot be reproduced from the submitted documentation.",
        ],
    )

    add_heading(doc, "12. Cost, safety and professional conduct", 1)
    add_list(
        doc,
        [
            "Use small development-sized resources and deploy the full stack only long enough to gather evidence.",
            "Create a low AWS Budget alert where account permissions allow it.",
            "Destroy NAT Gateways, load balancers, WAF resources and EC2 instances promptly after the demonstration.",
            "Do not use real customer data, production credentials or an employer-owned account without authorization.",
            "Never publish sensitive repository content to a public location.",
            "If a required service is unavailable in the account, document the limitation and provide deployable code rather than silently omitting it.",
        ],
    )
    add_callout(
        doc,
        "Cost expectation",
        "The assessment is designed for a short-lived sandbox deployment. Estimate cost before deployment and aim to keep actual spend low. Free-tier eligibility is not assumed.",
        fill="FFF8E8",
        accent="8A5A00",
    )

    add_heading(doc, "Appendix A. Reference data contract", 1)
    add_code_block(
        doc,
        """{
  "jobId": "550e8400-e29b-41d4-a716-446655440000",
  "tenantId": "tenant-001",
  "customerId": "CUST-123",
  "documentType": "invoice",
  "fileName": "invoice-2026-001.pdf",
  "amount": 1250.50,
  "createdAt": "2026-07-22T10:00:00Z"
}""",
    )

    add_heading(doc, "Appendix B. Suggested DynamoDB attributes", 1)
    add_table(
        doc,
        ["Attribute", "Purpose"],
        [
            ("tenantId", "Tenant partition boundary or part of the selected composite key."),
            ("jobId", "Immutable job identifier."),
            ("status", "PENDING, PROCESSING, COMPLETED, REJECTED or FAILED."),
            ("sourceObjectKey", "Tenant-scoped S3 intake key."),
            ("processedObjectKey", "Processed or quarantine key when present."),
            ("checksum", "SHA-256 checksum for accepted content."),
            ("createdAt / updatedAt", "UTC timestamps used for ordering and operational evidence."),
            ("failureReason", "Sanitized reason suitable for operators; no raw document body."),
            ("attemptCount", "Processing attempts or a justified equivalent."),
        ],
        [2600, 6760],
    )

    add_heading(doc, "Appendix C. Final submission checklist", 1)
    add_list(
        doc,
        [
            "All mandatory services are represented in code or an explicitly documented limitation.",
            "Deployment, verification and cleanup commands have been run from a clean environment where practical.",
            "No credentials, secrets, state files, build artifacts or sensitive values are committed.",
            "The architecture diagram matches the deployed solution.",
            "Successful, rejected, retry and DLQ paths have evidence.",
            "Security assumptions, costs, production gaps, RPO and RTO are documented.",
            "The repository is accessible to the hiring team for the agreed review period.",
        ],
    )

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
